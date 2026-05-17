"""
Cria o template DOCX da vistoria com colunas Entrada e Saída.
Roda uma vez para gerar VISTORIA_ENTRADA_SAIDA_TEMPLATE.docx.

Identidade visual: Ativuz Veículos
  - Logo no topo (assets/ativuz_logo.png)
  - Cor primária (cabeçalhos): #232938 (azul-quase-preto da marca)
  - Cor de destaque:           #FF5722 (laranja Ativuz)
  - Tudo centralizado horizontalmente
"""
from pathlib import Path

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ── Paleta da marca Ativuz ───────────────────────────────────────────────────
COR_MARCA_HEX   = '232938'   # azul-quase-preto (cabeçalhos)
COR_LARANJA_HEX = 'FF5722'   # laranja Ativuz (destaque)
COR_CINZA_HEX   = 'F2F2F2'   # cinza claro (labels)
COR_ZEBRA_HEX   = 'FAFAFA'   # zebra das linhas

COR_MARCA_RGB   = RGBColor(0x23, 0x29, 0x38)
COR_LARANJA_RGB = RGBColor(0xFF, 0x57, 0x22)
COR_CINZA_TXT   = RGBColor(0x66, 0x66, 0x66)
COR_BRANCO      = RGBColor(0xFF, 0xFF, 0xFF)

LOGO_PATH = Path(__file__).parent / "assets" / "ativuz_logo.png"

# ── Lista padrão de 30 acessórios (mantendo nomenclatura do template antigo) ──
ACESSORIOS = [
    ("acc_calotas",          "Calotas"),
    ("acc_buzina",           "Buzina"),
    ("acc_doc_crlv",         "DOC. CRLV"),
    ("acc_triangulo",        "Triângulo Sinaliz."),
    ("acc_antena",           "Antena"),
    ("acc_sensor_re",        "Sensor de Ré"),
    ("acc_som",              "Som / Alto-falante"),
    ("acc_tapetes",          "Tapetes"),
    ("acc_limpadores",       "Limpadores"),
    ("acc_chave_roda",       "Chave de Roda"),
    ("acc_vidros_eletricos", "Vidros Elétricos"),
    ("acc_oleo_motor",       "Óleo do Motor"),
    ("acc_alarme",           "Alarme / Travas"),
    ("acc_lampadas",         "Lâmpadas"),
    ("acc_macaco",           "Macaco Mecânico"),
    ("acc_estepe",           "Estepe"),
    ("acc_gnv",              "Func. GNV"),
    ("acc_agua",             "Água"),
    ("acc_borr_psg_dir",     "Borracha PSG Dir."),
    ("acc_borr_mtr_dir",     "Borracha MTR Dir."),
    ("acc_asa_dd",           "Asa Urubu D.D."),
    ("acc_asa_td",           "Asa Urubu T.D."),
    ("acc_tapete_mala",      "Tapete de Mala"),
    ("acc_tampa_parachoque", "Tampa Para-choque"),
    ("acc_borr_psg_tras",    "Borracha PSG Tras."),
    ("acc_borr_mtr_tras",    "Borracha MTR Tras."),
    ("acc_asa_de",           "Asa Urubu D.E."),
    ("acc_asa_te",           "Asa Urubu T.E."),
    ("acc_bagagito",         "Bagagito"),
    ("acc_lingueta",         "Lingueta"),
]

# ─────────────────────────────────────────────────────────────────────────────
def set_cell_bg(cell, hex_color):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

def add_borders(table, color='888888'):
    tbl = table._tbl
    tblPr = tbl.tblPr
    tblBorders = OxmlElement('w:tblBorders')
    for border in ('top','left','bottom','right','insideH','insideV'):
        b = OxmlElement(f'w:{border}')
        b.set(qn('w:val'), 'single')
        b.set(qn('w:sz'), '4')
        b.set(qn('w:color'), color)
        tblBorders.append(b)
    tblPr.append(tblBorders)

def cell(table, r, c, text, *, bold=False, size=10, bg=None, align='center', color=None):
    """Por padrão alinhamento CENTER — toda a vistoria fica centralizada."""
    tcell = table.cell(r, c)
    tcell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    tcell.text = ''
    p = tcell.paragraphs[0]
    p.alignment = {'left': WD_ALIGN_PARAGRAPH.LEFT,
                   'center': WD_ALIGN_PARAGRAPH.CENTER,
                   'right': WD_ALIGN_PARAGRAPH.RIGHT}[align]
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = color
    if bg:
        set_cell_bg(tcell, bg)
    return tcell

def add_section_title(doc, numero, titulo):
    """Cabeçalho de seção centralizado, com barra laranja sutil embaixo."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(f"{numero}. {titulo}")
    r.bold = True
    r.font.size = Pt(11)
    r.font.color.rgb = COR_MARCA_RGB
    return p

# ─────────────────────────────────────────────────────────────────────────────
doc = Document()

# Margens
for section in doc.sections:
    section.top_margin    = Cm(1.5)
    section.bottom_margin = Cm(1.5)
    section.left_margin   = Cm(1.5)
    section.right_margin  = Cm(1.5)

# ── Logo da marca (centralizada) ─────────────────────────────────────────────
if LOGO_PATH.exists():
    p_logo = doc.add_paragraph()
    p_logo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p_logo.add_run()
    run.add_picture(str(LOGO_PATH), width=Cm(7.0))
else:
    # Fallback: nome da empresa em texto se a logo não existir
    p_logo = doc.add_paragraph()
    p_logo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p_logo.add_run("ATIVUZ VEÍCULOS")
    run.bold = True
    run.font.size = Pt(20)
    run.font.color.rgb = COR_MARCA_RGB

# ── Cabeçalho ────────────────────────────────────────────────────────────────
h1 = doc.add_paragraph()
h1.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = h1.add_run("ANEXO I — CONTRATO DE LOCAÇÃO")
r.bold = True
r.font.size = Pt(12)
r.font.color.rgb = COR_MARCA_RGB

h2 = doc.add_paragraph()
h2.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = h2.add_run("VISTORIA DE ENTREGA E DEVOLUÇÃO DO VEÍCULO")
r.bold = True
r.font.size = Pt(13)
r.font.color.rgb = COR_LARANJA_RGB

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = sub.add_run("Contrato [contrato_id]   •   Documento único — preenchido em duas etapas")
r.italic = True
r.font.size = Pt(9)
r.font.color.rgb = COR_CINZA_TXT

doc.add_paragraph()  # espaço

# ── 1. Dados do cliente ──────────────────────────────────────────────────────
add_section_title(doc, "1", "DADOS DO CLIENTE")

t = doc.add_table(rows=3, cols=4)
t.alignment = WD_TABLE_ALIGNMENT.CENTER
add_borders(t)
cell(t, 0, 0, "Cliente",        bold=True, bg=COR_CINZA_HEX, align='center')
cell(t, 0, 1, "[cliente_nome]", align='center')
cell(t, 0, 2, "Telefone",       bold=True, bg=COR_CINZA_HEX, align='center')
cell(t, 0, 3, "[cliente_telefone]", align='center')
cell(t, 1, 0, "Endereço",       bold=True, bg=COR_CINZA_HEX, align='center')
t.cell(1, 1).merge(t.cell(1, 3))
cell(t, 1, 1, "[cliente_endereco]", align='center')
cell(t, 2, 0, "Preenchido por", bold=True, bg=COR_CINZA_HEX, align='center')
cell(t, 2, 1, "[preenchido_por]", align='center')
cell(t, 2, 2, "Contrato",       bold=True, bg=COR_CINZA_HEX, align='center')
cell(t, 2, 3, "[contrato_id]",  align='center')

doc.add_paragraph()

# ── 2. Dados do veículo ──────────────────────────────────────────────────────
add_section_title(doc, "2", "DADOS DO VEÍCULO")

t = doc.add_table(rows=3, cols=4)
t.alignment = WD_TABLE_ALIGNMENT.CENTER
add_borders(t)
cell(t, 0, 0, "Veículo",      bold=True, bg=COR_CINZA_HEX, align='center')
cell(t, 0, 1, "[veiculo]",    align='center')
cell(t, 0, 2, "Placa",        bold=True, bg=COR_CINZA_HEX, align='center')
cell(t, 0, 3, "[placa]",      align='center')
cell(t, 1, 0, "Cor",          bold=True, bg=COR_CINZA_HEX, align='center')
cell(t, 1, 1, "[cor]",        align='center')
cell(t, 1, 2, "Ano",          bold=True, bg=COR_CINZA_HEX, align='center')
cell(t, 1, 3, "[ano]",        align='center')
cell(t, 2, 0, "Chassi",       bold=True, bg=COR_CINZA_HEX, align='center')
cell(t, 2, 1, "[chassi]",     align='center')
cell(t, 2, 2, "Número Motor", bold=True, bg=COR_CINZA_HEX, align='center')
cell(t, 2, 3, "[numero_motor]", align='center')

doc.add_paragraph()

# ── 3. Comparativo Entrada × Saída (medidores) ───────────────────────────────
add_section_title(doc, "3", "ENTREGA × DEVOLUÇÃO")

t = doc.add_table(rows=4, cols=3)
t.alignment = WD_TABLE_ALIGNMENT.CENTER
add_borders(t)
# Cabeçalho com cor da marca
cell(t, 0, 0, "Item",      bold=True, bg=COR_MARCA_HEX, align='center', color=COR_BRANCO)
cell(t, 0, 1, "Entrega",   bold=True, bg=COR_MARCA_HEX, align='center', color=COR_BRANCO)
cell(t, 0, 2, "Devolução", bold=True, bg=COR_MARCA_HEX, align='center', color=COR_BRANCO)

cell(t, 1, 0, "Data e hora", bold=True, bg=COR_CINZA_HEX, align='center')
cell(t, 1, 1, "[data_entrada]",   align='center')
cell(t, 1, 2, "[data_saida]",     align='center')

cell(t, 2, 0, "Hodômetro (km)",   bold=True, bg=COR_CINZA_HEX, align='center')
cell(t, 2, 1, "[hodometro_entrada]", align='center')
cell(t, 2, 2, "[hodometro_saida]",   align='center')

cell(t, 3, 0, "Combustível",      bold=True, bg=COR_CINZA_HEX, align='center')
cell(t, 3, 1, "[combustivel_entrada]", align='center')
cell(t, 3, 2, "[combustivel_saida]",   align='center')

doc.add_paragraph()

# ── 4. Acessórios e equipamentos ─────────────────────────────────────────────
add_section_title(doc, "4", "ACESSÓRIOS E EQUIPAMENTOS")

leg = doc.add_paragraph()
leg.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = leg.add_run("S = Sim, existente     |     N = Não existente     |     A = Avariado")
r.font.size = Pt(8); r.italic = True
r.font.color.rgb = COR_CINZA_TXT

# Tabela com 4 colunas: Item | Entrega | Devolução | Status
t = doc.add_table(rows=len(ACESSORIOS) + 1, cols=4)
t.alignment = WD_TABLE_ALIGNMENT.CENTER
add_borders(t)
# largura aproximada
for col, w in zip(t.columns, [Cm(5.5), Cm(2.5), Cm(2.5), Cm(6.0)]):
    for c in col.cells:
        c.width = w

# Cabeçalho
cell(t, 0, 0, "Acessório",  bold=True, bg=COR_MARCA_HEX, align='center', color=COR_BRANCO)
cell(t, 0, 1, "Entrega",    bold=True, bg=COR_MARCA_HEX, align='center', color=COR_BRANCO)
cell(t, 0, 2, "Devolução",  bold=True, bg=COR_MARCA_HEX, align='center', color=COR_BRANCO)
cell(t, 0, 3, "Status / Observação", bold=True, bg=COR_MARCA_HEX, align='center', color=COR_BRANCO)

for i, (key, label) in enumerate(ACESSORIOS):
    row = i + 1
    bg = COR_ZEBRA_HEX if i % 2 == 0 else None
    cell(t, row, 0, label, size=9, bg=bg, align='center')
    cell(t, row, 1, f"[{key}_entrada]", size=9, align='center', bg=bg, bold=True)
    cell(t, row, 2, f"[{key}_saida]",   size=9, align='center', bg=bg, bold=True)
    cell(t, row, 3, f"[{key}_status]",  size=9, bg=bg, align='center')

doc.add_paragraph()

# ── 5. Observações ───────────────────────────────────────────────────────────
add_section_title(doc, "5", "OBSERVAÇÕES")

t = doc.add_table(rows=2, cols=2)
t.alignment = WD_TABLE_ALIGNMENT.CENTER
add_borders(t)
cell(t, 0, 0, "Observações — Entrega",   bold=True, bg=COR_CINZA_HEX, align='center')
cell(t, 0, 1, "Observações — Devolução", bold=True, bg=COR_CINZA_HEX, align='center')
cell(t, 1, 0, "[obs_entrada]", align='center')
cell(t, 1, 1, "[obs_saida]",   align='center')

doc.add_paragraph()

# ── 6. Sintomas / Danos ──────────────────────────────────────────────────────
add_section_title(doc, "6", "SINTOMAS, DANOS OU AVARIAS RELATADOS")

t = doc.add_table(rows=2, cols=2)
t.alignment = WD_TABLE_ALIGNMENT.CENTER
add_borders(t)
cell(t, 0, 0, "Na entrega",   bold=True, bg=COR_CINZA_HEX, align='center')
cell(t, 0, 1, "Na devolução", bold=True, bg=COR_CINZA_HEX, align='center')
cell(t, 1, 0, "[sintomas_entrada]", align='center')
cell(t, 1, 1, "[sintomas_saida]",   align='center')

doc.add_paragraph()

# ── 7. Assinaturas ───────────────────────────────────────────────────────────
add_section_title(doc, "7", "ASSINATURAS")

t = doc.add_table(rows=3, cols=2)
t.alignment = WD_TABLE_ALIGNMENT.CENTER
add_borders(t)
cell(t, 0, 0, "ENTREGA",   bold=True, bg=COR_MARCA_HEX, align='center', color=COR_BRANCO)
cell(t, 0, 1, "DEVOLUÇÃO", bold=True, bg=COR_MARCA_HEX, align='center', color=COR_BRANCO)
cell(t, 1, 0, "\n\n\n______________________________________\nAssinatura do Cliente\n[cliente_nome]", align='center', size=9)
cell(t, 1, 1, "\n\n\n______________________________________\nAssinatura do Cliente\n[cliente_nome]", align='center', size=9)
cell(t, 2, 0, "\n\n\n______________________________________\nResponsável Ativuz\n[responsavel_entrada]", align='center', size=9)
cell(t, 2, 1, "\n\n\n______________________________________\nResponsável Ativuz\n[responsavel_saida]",  align='center', size=9)

doc.add_paragraph()

# ── 8. Fotos ─────────────────────────────────────────────────────────────────
add_section_title(doc, "8", "REGISTRO FOTOGRÁFICO")

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("FOTOS DA ENTREGA")
r.bold = True; r.font.size = Pt(10)
r.font.color.rgb = COR_MARCA_RGB
p_marker = doc.add_paragraph("[FOTOS_ENTRADA]")
p_marker.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("FOTOS DA DEVOLUÇÃO")
r.bold = True; r.font.size = Pt(10)
r.font.color.rgb = COR_MARCA_RGB
p_marker = doc.add_paragraph("[FOTOS_SAIDA]")
p_marker.alignment = WD_ALIGN_PARAGRAPH.CENTER

# Salvar
out = Path(__file__).parent / "docx_templates" / "VISTORIA_ENTRADA_SAIDA_TEMPLATE.docx"
out.parent.mkdir(parents=True, exist_ok=True)
doc.save(str(out))
print(f"Template criado: {out}")
